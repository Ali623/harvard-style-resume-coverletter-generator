"""Build .docx documents from Resume and CoverLetter models."""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from resume_generator.renderer import (
    new_doc, set_page_letter, add_name_paragraph, add_contact_header,
    add_section_heading, add_bullet, make_invisible_table,
    set_line_spacing, add_run, FONT, HEADING_FONT, HEADER_FONT,
    BODY_PT, HEADING_PT, COMPANY_PT,
)
from resume_generator.models import Resume, CoverLetter
from resume_generator.validate import Budget


# ─── CHARACTER BUDGETS ───────────────────────────────────────────────────────

RESUME_CHARS = {
    "bullet":         90,
    "edu_uni":        50,
    "edu_location":   28,
    "edu_degree":     50,
    "edu_dates":      34,
    "exp_company":    50,
    "exp_location":   28,
    "exp_role":       50,
    "exp_dates":      28,
    "skill_label":    28,
    "skill_value":    72,
    "cert":           90,
    "achievement":    90,
    "body_line":     100,
}

COVER_CHARS = {
    "para":          600,
    "recipient_name": 50,
    "recipient_city": 50,
    "salutation":     50,
    "signoff":        40,
    "body_line":     105,
}


def _validate_resume(r: Resume) -> Budget:
    b = Budget()
    for edu in r.education:
        b.check(edu.uni,      RESUME_CHARS['edu_uni'],      'edu uni')
        b.check(edu.location, RESUME_CHARS['edu_location'], 'edu location')
        b.check(edu.degree,   RESUME_CHARS['edu_degree'],   'edu degree')
        b.check(edu.dates,    RESUME_CHARS['edu_dates'],    'edu dates')
    for exp in r.experience:
        b.check(exp.company,  RESUME_CHARS['exp_company'],  'exp company')
        b.check(exp.location, RESUME_CHARS['exp_location'], 'exp location')
        b.check(exp.role,     RESUME_CHARS['exp_role'],     'exp role')
        b.check(exp.dates,    RESUME_CHARS['exp_dates'],    'exp dates')
        for bl in exp.bullets:
            b.check(bl, RESUME_CHARS['bullet'], 'exp bullet')
    for sk in r.skills:
        b.check(sk.label, RESUME_CHARS['skill_label'], 'skill label')
        b.check(sk.value, RESUME_CHARS['skill_value'], 'skill value')
    for c in r.certifications:
        b.check(c, RESUME_CHARS['cert'], 'cert')
    for proj in r.projects:
        b.check(proj.title + proj.description, RESUME_CHARS['bullet'], 'project (title+desc)')
    for a in r.achievements:
        b.check(a, RESUME_CHARS['achievement'], 'achievement')
    return b


def _validate_cover(c: CoverLetter) -> Budget:
    b = Budget()
    b.check(c.recipient_name, COVER_CHARS['recipient_name'], 'recipient name')
    b.check(c.recipient_city, COVER_CHARS['recipient_city'], 'recipient city')
    b.check(c.salutation,     COVER_CHARS['salutation'],     'salutation')
    for i, p in enumerate(c.paragraphs, 1):
        b.check(p, COVER_CHARS['para'], f'paragraph {i}')
    return b


# ─── RESUME BUILDER ──────────────────────────────────────────────────────────

def build_resume(resume: Resume, output_dir: str = ".", company: str = "", job_title: str = "", timestamp: str = "") -> str:
    """Build resume .docx, validate char budgets, return path."""
    budget = _validate_resume(resume)
    filename = "Aliullah-Resume.docx"
    doc = new_doc()
    set_page_letter(doc)

    # Name
    add_name_paragraph(doc, resume.full_name)

    # Contact header
    add_contact_header(doc, resume.address, resume.phone, resume.email,
                       resume.linkedin, resume.github, resume.portfolio)

    # Professional Summary (if present)
    if resume.summary.strip():
        add_section_heading(doc, 'Professional Summary')
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_line_spacing(para)
        add_run(para, resume.summary, BODY_PT)

    # Education
    add_section_heading(doc, 'Education')
    for edu in resume.education:
        t = make_invisible_table(doc, 2, [6800, 4530])
        r0 = t.rows[0]
        r0.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_line_spacing(r0.cells[0].paragraphs[0])
        set_line_spacing(r0.cells[1].paragraphs[0])
        # bullet style on uni name
        pPr = r0.cells[0].paragraphs[0]._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); numPr.append(ilvl)
        numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1'); numPr.append(numId)
        pPr.append(numPr)
        add_run(r0.cells[0].paragraphs[0], edu.uni,      BODY_PT, bold=True)
        add_run(r0.cells[1].paragraphs[0], edu.location, BODY_PT)
        r1 = t.rows[1]
        r1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_line_spacing(r1.cells[0].paragraphs[0])
        set_line_spacing(r1.cells[1].paragraphs[0])
        add_run(r1.cells[0].paragraphs[0], edu.degree, BODY_PT, bold=True, italic=True)
        add_run(r1.cells[0].paragraphs[0], edu.grade,  BODY_PT)
        add_run(r1.cells[1].paragraphs[0], edu.dates,  BODY_PT, italic=True)

    # Experience
    add_section_heading(doc, 'Experience')
    for exp in resume.experience:
        ct = make_invisible_table(doc, 1, [6800, 4530])
        cr = ct.rows[0]
        cr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cr.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_line_spacing(cr.cells[0].paragraphs[0], 240)
        set_line_spacing(cr.cells[1].paragraphs[0], 240)
        add_run(cr.cells[0].paragraphs[0], exp.company,  COMPANY_PT, bold=True)
        add_run(cr.cells[1].paragraphs[0], exp.location, COMPANY_PT)

        rt = make_invisible_table(doc, 1, [6800, 4530])
        rr = rt.rows[0]
        rr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_line_spacing(rr.cells[0].paragraphs[0], 240)
        set_line_spacing(rr.cells[1].paragraphs[0], 240)
        add_run(rr.cells[0].paragraphs[0], exp.role,  BODY_PT, italic=True)
        add_run(rr.cells[1].paragraphs[0], exp.dates, BODY_PT, italic=True)

        for b_text in exp.bullets:
            add_bullet(doc, b_text)

    # Digital Competencies
    add_section_heading(doc, 'Digital Competencies')
    st = make_invisible_table(doc, len(resume.skills), [3200, 8130])
    for i, sk in enumerate(resume.skills):
        row = st.rows[i]
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_line_spacing(row.cells[0].paragraphs[0])
        set_line_spacing(row.cells[1].paragraphs[0])
        add_run(row.cells[0].paragraphs[0], sk.label, BODY_PT, bold=True)
        add_run(row.cells[1].paragraphs[0], sk.value, BODY_PT)

    # Certifications
    add_section_heading(doc, 'Certifications')
    for c in resume.certifications:
        add_bullet(doc, c)

    # Projects
    add_section_heading(doc, 'Projects')
    for proj in resume.projects:
        para = doc.add_paragraph(style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_line_spacing(para)
        add_run(para, proj.title, BODY_PT, bold=True)
        add_run(para, proj.description, BODY_PT)

    # Achievements
    add_section_heading(doc, 'Achievements')
    for a in resume.achievements:
        add_bullet(doc, a)

    # Validate & save
    budget.report()

    import os
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, filename)
    doc.save(path)
    print(f"Written: {path}")
    return path


# ─── COVER LETTER BUILDER ────────────────────────────────────────────────────

def build_coverletter(cover: CoverLetter, output_dir: str = ".", company: str = "", job_title: str = "", timestamp: str = "") -> str:
    """Build cover letter .docx, validate char budgets, return path."""
    budget = _validate_cover(cover)
    filename = "Aliullah-CoverLetter.docx"
    doc = new_doc()
    set_page_letter(doc, top=720, right=540, bottom=360, left=540)

    # Name
    add_name_paragraph(doc, cover.full_name)

    # Contact header
    add_contact_header(doc, cover.address, cover.phone, cover.email,
                       cover.linkedin, cover.github, cover.portfolio)

    # Recipient
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p)
    add_run(p, cover.recipient_name, BODY_PT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p)
    add_run(p, cover.recipient_city, BODY_PT)

    # Salutation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p)
    add_run(p, cover.salutation, BODY_PT)

    # Body paragraphs
    for para_text in cover.paragraphs:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_line_spacing(para)
        add_run(para, para_text, BODY_PT)

    # Sign-off
    for text, bold in [
        ('Best regards,', False),
        (cover.signoff_name, True),
        (cover.signoff_email, False),
        (cover.signoff_phone, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_line_spacing(p)
        add_run(p, text, BODY_PT, bold=bold)

    # Validate & save
    budget.report()

    import os
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, filename)
    doc.save(path)
    print(f"Written: {path}")
    return path
