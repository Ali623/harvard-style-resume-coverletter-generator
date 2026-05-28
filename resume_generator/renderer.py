"""Shared docx rendering helpers — invisible tables, borders, spacing, runs."""

from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Design tokens (can be overridden per document)
FONT         = 'Calibri'
HEADING_FONT = 'Times New Roman'
HEADER_FONT  = 'Calibri'
BODY_PT      = Pt(11)
HEADING_PT   = Pt(14)
COMPANY_PT   = Pt(11)
NAME_PT      = Pt(22)
LINE_SPACING = 240  # DXA


def set_cell_border_white(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'FFFFFF')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_border_white(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def set_col_width(table, col_widths_dxa):
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    else:
        for child in list(tblGrid):
            tblGrid.remove(child)
    for w in col_widths_dxa:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(col_widths_dxa[i % len(col_widths_dxa)]))
            tcW.set(qn('w:type'), 'dxa')


def set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_line_spacing(para, spacing_dxa=240):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(spacing_dxa))
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')


def add_run(para, text, font_size, bold=False, italic=False, font_name=FONT):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    return run


def make_invisible_table(doc, rows, col_widths):
    table = doc.add_table(rows=rows, cols=len(col_widths))
    set_table_border_white(table)
    set_col_width(table, col_widths)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border_white(cell)
            set_cell_margins_zero(cell)
    return table


def set_page_letter(doc, top=288, right=360, bottom=288, left=360):
    """US Letter: 8.5 x 11 inches."""
    section = doc.sections[0]
    section.page_width = Emu(12240 * 635)
    section.page_height = Emu(15840 * 635)
    section.top_margin = Emu(top * 635)
    section.right_margin = Emu(right * 635)
    section.bottom_margin = Emu(bottom * 635)
    section.left_margin = Emu(left * 635)


def add_section_heading(doc, text):
    """Section heading: HEADING_FONT, HEADING_PT, bold, centered, with bottom border."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.get_or_add_pPr()
    # bottom border
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    # spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '80')
    spacing.set(qn('w:after'), '50')
    pPr.append(spacing)
    # keep with next
    keep = OxmlElement('w:keepNext')
    pPr.append(keep)
    add_run(para, text, HEADING_PT, bold=True, font_name=HEADING_FONT)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(para)
    add_run(para, text, BODY_PT)
    return para


def new_doc():
    """Create a fresh Document with zero default paragraph spacing."""
    from docx import Document
    doc = Document()
    for style in doc.styles:
        try:
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
    return doc


def add_name_paragraph(doc, name):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '30')
    pPr.append(sp)
    add_run(para, name, NAME_PT, bold=True, font_name=HEADING_FONT)
    return para


def add_hyperlink(para, text: str, url: str, font_size, font_name=FONT):
    """Add a clickable hyperlink run to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # half-points
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    run = OxmlElement('w:r')
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    hl.append(run)
    para._p.append(hl)


def add_contact_header(doc, address, phone, email, linkedin, github, portfolio):
    """3-column, 2-row invisible contact table with clickable links."""
    COL3 = [3776, 3777, 3777]
    table = make_invisible_table(doc, 2, COL3)

    row0 = table.rows[0]
    row1 = table.rows[1]
    for row in (row0, row1):
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(cell.paragraphs[0], 240)

    add_run(row0.cells[0].paragraphs[0], address, BODY_PT, font_name=HEADER_FONT)
    add_run(row0.cells[1].paragraphs[0], phone,   BODY_PT, font_name=HEADER_FONT)
    add_hyperlink(row0.cells[2].paragraphs[0], email, f"mailto:{email}", BODY_PT, HEADER_FONT)

    add_hyperlink(row1.cells[0].paragraphs[0], linkedin,  f"https://{linkedin}",  BODY_PT, HEADER_FONT)
    add_hyperlink(row1.cells[1].paragraphs[0], github,    f"https://{github}",    BODY_PT, HEADER_FONT)
    add_hyperlink(row1.cells[2].paragraphs[0], portfolio, f"https://{portfolio}", BODY_PT, HEADER_FONT)

    return table
